import streamlit as st
from pypdf import PdfReader
from docx import Document
from ai import ask_ai
from io import BytesIO


# ---------------- PDF TEXT EXTRACTION ----------------

def extract_text_from_pdf(uploaded_file):

    text = ""

    try:
        reader = PdfReader(uploaded_file)

        for page in reader.pages:
            page_text = page.extract_text()

            if page_text:
                text += page_text + "\n"

    except Exception as e:
        st.error(f"Error reading PDF: {e}")

    return text


# ---------------- DOCX TEXT EXTRACTION ----------------

def extract_text_from_docx(uploaded_file):

    text = ""

    try:
        doc = Document(uploaded_file)

        for para in doc.paragraphs:
            text += para.text + "\n"

    except Exception as e:
        st.error(f"Error reading DOCX: {e}")

    return text


# ---------------- CREATE DOCX FILE ----------------

def create_docx(summary_text):

    doc = Document()

    doc.add_heading("Lesson Summary", level=1)

    for line in summary_text.split("\n"):
        doc.add_paragraph(line)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer


# ---------------- MAIN FUNCTION ----------------

def summarize():

    st.subheader("📚 Summarize Your Lesson")

    st.write("Upload a lesson file or paste lesson text.")

    input_method = st.radio(
        "Choose Input Method",
        ["Upload File", "Paste Lesson Text"]
    )

    text = ""


    # ---------- FILE UPLOAD ----------

    if input_method == "Upload File":

        uploaded_file = st.file_uploader(
            "Upload Lesson File",
            type=["pdf", "docx"]
        )

        if uploaded_file:

            file_type = uploaded_file.name.split(".")[-1]

            with st.spinner("📖 Reading lesson file..."):

                if file_type == "pdf":
                    text = extract_text_from_pdf(uploaded_file)

                elif file_type == "docx":
                    text = extract_text_from_docx(uploaded_file)


    # ---------- TEXT INPUT ----------

    else:

        text = st.text_area(
            "Paste Lesson Content",
            height=300,
            placeholder="Paste lesson notes, lecture material, or teaching content here..."
        )


    if not text:
        return


    word_count = len(text.split())

    st.info(f"Lesson contains approximately **{word_count} words**.")


    with st.expander("Preview Lesson Text"):
        st.write(text[:2000])


    # ---------- GENERATE SUMMARY ----------

    if st.button("Generate Lesson Summary"):

        MAX_TEXT = 40000
        text = text[:MAX_TEXT]

        system_prompt = """
You are an expert academic assistant helping teachers.

Summarize lessons clearly using structured bullet points.

Focus on:
- Key Concepts
- Important Definitions
- Core Ideas
- Practical Examples
- Teaching Notes

Keep summaries concise and classroom friendly.
"""

        user_prompt = f"""
Lesson Content:
{text}

Create a clear lesson summary for teachers using bullet points.
"""

        with st.spinner("🤖 AI is analyzing the lesson and generating summary..."):

            summary = ask_ai(system_prompt, user_prompt)


        st.success("Summary generated successfully!")

        st.session_state["summary"] = summary


    # ---------- DISPLAY SUMMARY ----------

    if "summary" in st.session_state:

        summary = st.session_state["summary"]

        st.markdown("## 📘 Lesson Summary")

        st.markdown(
            f"""
<div style="
border:1px solid #ddd;
padding:20px;
border-radius:10px;
font-size:16px;
line-height:1.6;
color:#000;
background-color:#fafafa;
">
{summary}
</div>
""",
            unsafe_allow_html=True
        )


        # ---------- CREATE DOCX ----------

        docx_file = create_docx(summary)


        # ---------- DOWNLOAD BUTTON ----------

        downloaded = st.download_button(
            label="⬇ Download Summary (DOCX)",
            data=docx_file,
            file_name="lesson_summary.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )


        # ---------- RESET AFTER DOWNLOAD ----------

        if downloaded:
            st.session_state.clear()
            st.rerun()