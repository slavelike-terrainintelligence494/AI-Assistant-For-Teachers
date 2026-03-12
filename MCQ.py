import streamlit as st
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from ai import ask_ai
from concurrent.futures import ThreadPoolExecutor, as_completed


# ---------------- GENERATE ONE MCQ ----------------
def generate_single_mcq(topic, difficulty, q_number):

    system_prompt = """
You are an expert educational assessment designer.
Generate high-quality multiple-choice questions.
"""

    user_prompt = f"""
Generate ONE multiple-choice question.

Topic: {topic}
Difficulty: {difficulty}

Rules:
- 4 options only (a,b,c,d)
- Options must be short (2–3 words)
- Include correct answer
- No coding questions

Format EXACTLY:

Q{q_number}: Question text
a. Option
b. Option
c. Option
d. Option
Answer: a
"""

    return ask_ai(system_prompt, user_prompt)


# ---------------- FORMAT QUIZ ----------------
def format_quiz(quiz_text):

    lines = quiz_text.split("\n")

    formatted_quiz = []
    current_question = []

    for line in lines:

        if re.match(r"^Q\d+:", line):
            if current_question:
                formatted_quiz.append(current_question)
            current_question = [line]

        elif re.match(r"^[a-d]\.", line):
            current_question.append(line)

        elif line.startswith("Answer:"):
            current_question.append(line)

    if current_question:
        formatted_quiz.append(current_question)

    return formatted_quiz


# ---------------- BALANCED DIFFICULTY ----------------
def create_difficulty_mix(total):

    beginner = total // 3
    intermediate = total // 3
    expert = total - beginner - intermediate

    difficulties = (
        ["Beginner"] * beginner
        + ["Intermediate"] * intermediate
        + ["Expert"] * expert
    )

    return difficulties


# ---------------- PARALLEL QUIZ GENERATION ----------------
def generate_quiz_parallel(topic, difficulty, num_questions, balanced):

    results = []

    if balanced:
        difficulty_list = create_difficulty_mix(num_questions)
    else:
        difficulty_list = [difficulty] * num_questions

    with ThreadPoolExecutor(max_workers=5) as executor:

        futures = []

        for i, diff in enumerate(difficulty_list, start=1):
            futures.append(
                executor.submit(generate_single_mcq, topic, diff, i)
            )

        for future in as_completed(futures):
            results.append(future.result())

    return results


# ---------------- CREATE DOCX ----------------
def generate_docx(quiz, heading1, heading2):

    doc = Document()

    h1 = doc.add_heading(heading1, level=0)
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    h2 = doc.add_heading(heading2, level=2)
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("Name:")
    doc.add_paragraph("Student ID:")
    doc.add_paragraph("Class:")
    doc.add_paragraph("Section:")
    doc.add_paragraph("")

    # Questions
    for question in quiz:

        for line in question:
            if not line.startswith("Answer:"):
                doc.add_paragraph(line)

        doc.add_paragraph("")

    # Answer Key
    doc.add_page_break()
    doc.add_heading("Answer Key", level=1)

    for i, question in enumerate(quiz):

        for line in question:
            if line.startswith("Answer:"):
                ans = line.split(": ")[1]
                doc.add_paragraph(f"Q{i+1}: {ans}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer


# ---------------- STREAMLIT APP ----------------
def MCQ():

    st.subheader("AI MCQ Quiz Generator")

    col1, col2 = st.columns(2)

    with col1:
        institute_name = st.text_input("Institute Name")
        topic = st.text_input("Topic")

    with col2:
        quiz_title = st.text_input("Quiz Title")

        difficulty = st.selectbox(
            "Difficulty Level",
            ["Beginner", "Intermediate", "Expert"]
        )

    num_questions = st.number_input(
        "Number of Questions",
        min_value=1,
        max_value=20,
        value=5
    )

    balanced = st.checkbox(
        "Enable AI Difficulty Balancing (mixed difficulty)",
        value=False
    )

    show_answers = st.checkbox("Show Answers", value=True)

    # ---------- GENERATE QUIZ ----------
    if st.button("Generate Quiz"):

        if topic.strip() == "":
            st.error("Please enter a topic.")
            return

        with st.spinner("Generating quiz using AI (parallel mode)..."):

            raw_questions = generate_quiz_parallel(
                topic,
                difficulty,
                num_questions,
                balanced
            )

        formatted_quiz = []

        for q in raw_questions:
            formatted_quiz.extend(format_quiz(q))

        st.session_state["quiz"] = formatted_quiz

        docx_content = generate_docx(
            formatted_quiz,
            institute_name or "Institute",
            quiz_title or "Quiz"
        )

        st.session_state["docx_content"] = docx_content

    # ---------- DISPLAY QUIZ ----------
    if "quiz" in st.session_state:

        st.divider()

        if institute_name:
            st.write(f"### {institute_name}")

        if quiz_title:
            st.write(f"**{quiz_title}**")

        st.write("---")

        st.write("Name:")
        st.write("Student ID:")
        st.write("Class:")
        st.write("Section:")
        st.write("")

        for question in st.session_state["quiz"]:

            st.markdown(f"**{question[0]}**")

            for line in question[1:]:

                if line.startswith("Answer:"):

                    if show_answers:
                        st.success(f"Correct Answer: {line.split(': ')[1]}")

                else:
                    st.write(line)

            st.write("")

        # ---------- ANSWER KEY ----------
        if show_answers:

            st.subheader("Answer Key")

            for i, question in enumerate(st.session_state["quiz"]):

                for line in question:
                    if line.startswith("Answer:"):
                        ans = line.split(": ")[1]
                        st.write(f"Q{i+1}: {ans}")

        # ---------- DOWNLOAD ----------
        if "docx_content" in st.session_state:

            st.download_button(
                label="Download Quiz as DOCX",
                data=st.session_state["docx_content"],
                file_name="quiz.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )