import streamlit as st
from ai import ask_ai
from docx import Document
from io import BytesIO


# ---------------- CREATE DOCX ----------------

def create_docx(chat_history):

    doc = Document()
    doc.add_heading("AI Counselling Conversation", level=1)

    for msg in chat_history:
        role = msg["role"].capitalize()
        content = msg["content"]
        doc.add_paragraph(f"{role}: {content}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer


# ---------------- COUNSELLOR ----------------

def counsellor():

    st.subheader("🧠 AI Counsellor for Teacher Wellness")

    st.markdown(
        "Chat with our AI Counsellor to seek support for stress, burnout, or classroom challenges."
    )

    st.info(
        "⚠️ This AI counsellor provides supportive guidance but is not a replacement for professional mental health care."
    )

    # ---------------- RESET CHAT ----------------

    if st.button("Start New Conversation"):
        st.session_state.messages = []
        st.rerun()


    # ---------------- QUICK TOPICS ----------------

    st.markdown("### Common Topics")

    col1, col2, col3 = st.columns(3)

    if col1.button("Stress"):
        st.session_state.messages.append(
            {"role": "user", "content": "I feel stressed from my teaching workload."}
        )
        st.rerun()

    if col2.button("Burnout"):
        st.session_state.messages.append(
            {"role": "user", "content": "I feel burned out from teaching responsibilities."}
        )
        st.rerun()

    if col3.button("Classroom Challenges"):
        st.session_state.messages.append(
            {"role": "user", "content": "I'm struggling with difficult classroom behaviour."}
        )
        st.rerun()


    # ---------------- INIT CHAT ----------------

    if "messages" not in st.session_state:
        st.session_state.messages = []


    # ---------------- DISPLAY CHAT ----------------

    for message in st.session_state.messages:

        with st.chat_message(message["role"]):
            st.markdown(message["content"])


    st.divider()


    # ---------------- USER INPUT ----------------

    if prompt := st.chat_input("How may I support you today?"):

        st.session_state.messages.append(
            {"role": "user", "content": prompt}
        )

        with st.chat_message("user"):
            st.markdown(prompt)


        # ---------------- AI RESPONSE ----------------

        with st.chat_message("assistant"):

            with st.spinner("AI counsellor is thinking..."):

                MAX_HISTORY = 10

                recent_messages = st.session_state.messages[-MAX_HISTORY:]

                conversation = ""

                for msg in recent_messages:
                    role = msg["role"]
                    content = msg["content"]
                    conversation += f"{role}: {content}\n"


                system_prompt = """
You are a warm and compassionate AI counsellor supporting teachers.

Speak like a supportive human conversation partner.

Guidelines:
- Start with empathy
- Use natural conversational language
- Keep responses short and friendly
- Give only 1–2 practical suggestions
- Ask thoughtful follow-up questions

Tone:
- calm
- supportive
- friendly
- conversational

Do NOT:
- give medical diagnoses
- sound like an academic essay
- give long lectures

Structure responses like a real conversation:
1. Empathy
2. Small helpful idea
3. Gentle question
"""


                user_prompt = f"""
Conversation so far:
{conversation}

Respond to the latest teacher message in a natural, human-like way.
"""


                response = ask_ai(system_prompt, user_prompt)

            st.markdown(response)


        st.session_state.messages.append(
            {"role": "assistant", "content": response}
        )


    # ---------------- DOWNLOAD DOCX ----------------

    if st.session_state.messages:

        docx_file = create_docx(st.session_state.messages)

        st.download_button(
            label="⬇ Download Conversation (DOCX)",
            data=docx_file,
            file_name="teacher_wellness_chat.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )